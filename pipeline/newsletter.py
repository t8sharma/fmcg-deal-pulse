"""
Newsletter builder.

Turns scored, de-duplicated, relevant articles into a short structured briefing a
business user can skim. Produces:
  * a Markdown string (used by the Streamlit app and as a portable artifact), and
  * a Word (.docx) document (the primary deliverable).

The structure is intentionally scannable: headline stats -> top deals (ranked) ->
category breakdown -> methodology & assumptions.
"""
from __future__ import annotations

import datetime as dt
from collections import Counter

from . import config


def _deal_line(a: dict) -> str:
    parties = ""
    if a.get("acquirer") and a.get("target"):
        parties = f"**{a['acquirer']} → {a['target']}**. "
    value = f" ({a['deal_value']})" if a.get("deal_value") else ""
    corro = ""
    if a.get("dup_count", 1) > 1:
        corro = f" · corroborated by {a['dup_count']} sources"
    return (
        f"{parties}{a.get('title_clean', a.get('title',''))}{value}\n"
        f"  _Source: {a.get('publisher') or a.get('domain') or 'n/a'} · "
        f"credibility {a.get('credibility',0):.2f} · relevance {a.get('relevance_score',0)}"
        f"{corro}_  \n  {a.get('url','')}"
    )


def build_markdown(articles: list[dict], generated_at: dt.datetime | None = None,
                   stats: dict | None = None) -> str:
    generated_at = generated_at or dt.datetime.now()
    stats = stats or {}
    ranked = sorted(
        articles,
        key=lambda x: (x.get("relevance_score", 0), x.get("credibility", 0), x.get("dup_count", 1)),
        reverse=True,
    )
    top = ranked[: config.TOP_DEALS_COUNT]

    cats = Counter(a.get("category", "Other / Diversified") for a in articles)
    lines: list[str] = []
    lines.append(f"# {config.NEWSLETTER_TITLE}")
    lines.append(f"### {config.NEWSLETTER_SUBTITLE}")
    lines.append(f"_Generated {generated_at.strftime('%d %b %Y, %H:%M')}_\n")

    # Headline stats
    lines.append("## At a glance")
    lines.append(
        f"- **{len(articles)}** relevant FMCG deals tracked "
        f"from **{stats.get('sources_count', len(set(a.get('publisher') for a in articles)))}** sources"
    )
    if stats:
        lines.append(
            f"- Pipeline funnel: {stats.get('ingested','?')} ingested → "
            f"{stats.get('after_dedup','?')} after de-duplication "
            f"({stats.get('duplicates_removed','?')} removed) → "
            f"{stats.get('relevant','?')} relevant"
        )
    top_cat = ", ".join(f"{c} ({n})" for c, n in cats.most_common(3))
    if top_cat:
        lines.append(f"- Most active categories: {top_cat}")
    lines.append("")

    # Top deals
    lines.append("## Top deals this cycle")
    for i, a in enumerate(top, 1):
        lines.append(f"{i}. {_deal_line(a)}\n")

    # Category breakdown
    lines.append("## By category")
    for cat, n in cats.most_common():
        lines.append(f"- **{cat}**: {n} deal(s)")
    lines.append("")

    # Methodology
    lines.append("## How this was built (methodology & assumptions)")
    lines.append(
        "- **Ingestion:** keyless public news (Google News RSS) across FMCG + deal queries.\n"
        "- **De-duplication:** exact URL/title match, then near-duplicate clustering "
        "(token-Jaccard + sequence similarity ≥ "
        f"{config.NEAR_DUP_THRESHOLD}); the highest-credibility copy is kept and the "
        "rest counted as corroboration.\n"
        "- **Relevance:** an article must carry BOTH a deal signal and an FMCG signal; "
        f"scored 0–100 and kept above {config.RELEVANCE_THRESHOLD}.\n"
        "- **Credibility:** transparent source tiers (wire/financial = 1.0, trade = 0.75, "
        "general = 0.5, unknown = 0.4); newsletter includes sources at or above "
        f"{config.CREDIBILITY_FLOOR}.\n"
        "- **Assumptions:** headlines/summaries are accurate; recency window = "
        f"{config.RECENCY_DAYS} days; deal value/parties are best-effort regex extractions "
        "and should be verified against the linked source before any decision."
    )
    lines.append("\n---\n_Auto-generated draft — review before distribution._")
    return "\n".join(lines)


def build_docx(articles: list[dict], out_path: str,
               generated_at: dt.datetime | None = None, stats: dict | None = None) -> str:
    """Write a formatted Word newsletter. Requires python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    generated_at = generated_at or dt.datetime.now()
    stats = stats or {}
    ranked = sorted(
        articles,
        key=lambda x: (x.get("relevance_score", 0), x.get("credibility", 0), x.get("dup_count", 1)),
        reverse=True,
    )
    top = ranked[: config.TOP_DEALS_COUNT]
    cats = Counter(a.get("category", "Other / Diversified") for a in articles)

    ACCENT = RGBColor(0x0B, 0x5C, 0x4A)  # deep green
    GREY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Masthead
    h = doc.add_paragraph()
    r = h.add_run(config.NEWSLETTER_TITLE)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    rs = sub.add_run(config.NEWSLETTER_SUBTITLE)
    rs.italic = True
    rs.font.size = Pt(12)
    rs.font.color.rgb = GREY
    meta = doc.add_paragraph()
    rm = meta.add_run(f"Generated {generated_at.strftime('%d %b %Y, %H:%M')}  ·  "
                      f"{len(articles)} deals  ·  "
                      f"{stats.get('sources_count','?')} sources")
    rm.font.size = Pt(9)
    rm.font.color.rgb = GREY

    doc.add_paragraph()  # spacer

    # At a glance
    doc.add_heading("At a glance", level=1)
    g = doc.add_paragraph(style="List Bullet")
    g.add_run(
        f"Pipeline funnel: {stats.get('ingested','?')} ingested → "
        f"{stats.get('after_dedup','?')} after de-duplication "
        f"({stats.get('duplicates_removed','?')} duplicates removed) → "
        f"{stats.get('relevant','?')} relevant deals."
    )
    if cats:
        top_cat = ", ".join(f"{c} ({n})" for c, n in cats.most_common(3))
        doc.add_paragraph(f"Most active categories: {top_cat}.", style="List Bullet")
    avg_cred = sum(a.get("credibility", 0) for a in articles) / max(len(articles), 1)
    doc.add_paragraph(f"Average source credibility: {avg_cred:.2f} (0–1 scale).",
                      style="List Bullet")

    # Top deals
    doc.add_heading("Top deals this cycle", level=1)
    for i, a in enumerate(top, 1):
        p = doc.add_paragraph()
        head = p.add_run(f"{i}. ")
        head.bold = True
        if a.get("acquirer") and a.get("target"):
            party = p.add_run(f"{a['acquirer']} → {a['target']}: ")
            party.bold = True
            party.font.color.rgb = ACCENT
        title_run = p.add_run(a.get("title_clean", a.get("title", "")))
        if a.get("deal_value"):
            vr = p.add_run(f"  [{a['deal_value']}]")
            vr.bold = True
        # meta line
        m = doc.add_paragraph()
        corro = f" · corroborated by {a['dup_count']} sources" if a.get("dup_count", 1) > 1 else ""
        mr = m.add_run(
            f"{a.get('category','')}  ·  Source: {a.get('publisher') or a.get('domain') or 'n/a'}  ·  "
            f"credibility {a.get('credibility',0):.2f}  ·  relevance {a.get('relevance_score',0)}{corro}"
        )
        mr.font.size = Pt(8.5)
        mr.font.color.rgb = GREY
        if a.get("url"):
            lk = doc.add_paragraph()
            lr = lk.add_run(a["url"])
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = RGBColor(0x1A, 0x5F, 0xB4)
            lr.italic = True

    # Category breakdown table
    doc.add_heading("By category", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Deals"
    for cat, n in cats.most_common():
        row = table.add_row().cells
        row[0].text = cat
        row[1].text = str(n)

    # Methodology
    doc.add_heading("Methodology & assumptions", level=1)
    method = [
        "Ingestion: keyless public news (Google News RSS) across FMCG + deal-language queries.",
        f"De-duplication: exact URL/title match, then near-duplicate clustering "
        f"(token-Jaccard + sequence similarity ≥ {config.NEAR_DUP_THRESHOLD}); the highest-"
        "credibility copy is kept, others counted as corroboration.",
        "Relevance: an article must carry BOTH a deal signal AND an FMCG signal; scored "
        f"0–100 and kept above {config.RELEVANCE_THRESHOLD}.",
        "Credibility: transparent source tiers (wire/financial 1.0, trade 0.75, general 0.5, "
        f"unknown 0.4); deals shown are at or above {config.CREDIBILITY_FLOOR}.",
        f"Assumptions: recency window = {config.RECENCY_DAYS} days; deal value and parties are "
        "best-effort extractions and should be verified against the linked source.",
    ]
    for m in method:
        doc.add_paragraph(m, style="List Bullet")

    foot = doc.add_paragraph()
    fr = foot.add_run("Auto-generated draft — review before distribution.")
    fr.italic = True
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = GREY

    doc.save(out_path)
    return out_path
